"""
Модуль для обработки DOCX файлов и извлечения текста с HTML разметкой
"""

import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO


def process_docx_file(file_content):
    """
    Обрабатывает DOCX файл и извлекает текст с HTML разметкой
    
    Args:
        file_content (bytes): Содержимое DOCX файла в байтах
        
    Returns:
        str: HTML разметка документа
        
    Raises:
        Exception: При ошибке обработки файла
    """
    try:
        file_stream = BytesIO(file_content)
        document = Document(file_stream)
        
        # Анализируем документ для определения размеров шрифтов
        font_sizes = analyze_font_sizes(document)
        base_font_size = font_sizes.get('base', 12)
        
        html_elements = []
        paragraph_counter = 1
        h1_used = False
        
        # Обрабатываем параграфы
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                cleaned_text = clean_text(text)
                if cleaned_text:
                    element = process_paragraph(
                        paragraph, cleaned_text, base_font_size, 
                        paragraph_counter, h1_used
                    )
                    
                    if element:
                        html_elements.append(element)
                        # Если это заголовок h1, отмечаем что он уже использован
                        if element.startswith('<h1>'):
                            h1_used = True
                        # Увеличиваем счетчик только для обычных абзацев
                        elif element.startswith('<p'):
                            paragraph_counter += 1
        
        # Обрабатываем таблицы
        table_html = extract_tables_as_html(document)
        if table_html:
            html_elements.append(table_html)
        
        return ''.join(html_elements)
        
    except Exception as e:
        raise Exception(f"Ошибка при обработке DOCX файла: {str(e)}")


def analyze_font_sizes(document):
    """
    Анализирует размеры шрифтов в документе для определения базового размера
    
    Args:
        document: Объект документа python-docx
        
    Returns:
        dict: Статистика размеров шрифтов
    """
    font_sizes = []
    
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.size:
                # Размер в пунктах
                size_pt = run.font.size.pt
                font_sizes.append(size_pt)
    
    if font_sizes:
        # Находим наиболее часто используемый размер как базовый
        size_counts = {}
        for size in font_sizes:
            size_counts[size] = size_counts.get(size, 0) + 1
        
        base_size = max(size_counts.items(), key=lambda x: x[1])[0]
        return {
            'base': base_size,
            'min': min(font_sizes),
            'max': max(font_sizes),
            'all': font_sizes
        }
    
    return {'base': 12, 'min': 12, 'max': 12, 'all': []}


def is_heading(paragraph, text, base_font_size):
    """
    Определяет, является ли параграф заголовком
    
    Args:
        paragraph: Объект параграфа python-docx
        text (str): Текст параграфа
        base_font_size (float): Базовый размер шрифта
        
    Returns:
        int: Уровень заголовка (1, 2, 3) или 0 если не заголовок
    """
    # Проверяем длину текста (заголовки обычно короткие)
    if len(text) > 200:  # Длинный текст вряд ли заголовок
        return 0
    
    # Проверяем количество слов
    word_count = len(text.split())
    if word_count > 25:  # Слишком много слов для заголовка
        return 0
    
    is_bold = False
    is_large_font = False
    is_centered = False
    max_font_size = base_font_size
    
    # Проверяем выравнивание
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        is_centered = True
    
    # Проверяем форматирование текста
    for run in paragraph.runs:
        if run.bold:
            is_bold = True
        
        if run.font.size:
            font_size = run.font.size.pt
            max_font_size = max(max_font_size, font_size)
            
            # Если шрифт больше базового на 2+ пункта
            if font_size > base_font_size + 1:
                is_large_font = True
    
    # Определяем уровень заголовка
    score = 0
    if is_bold:
        score += 3
    if is_large_font:
        score += 2
    if is_centered:
        score += 5
    if word_count <= 10:
        score += 1
    
    # Определяем уровень на основе размера шрифта и других факторов
    if max_font_size > base_font_size + 4:  # Очень большой шрифт
        return 1 if score >= 2 else 0
    elif max_font_size > base_font_size + 2:  # Большой шрифт
        return 2 if score >= 2 else 0
    elif score >= 3:  # Достаточно других признаков
        return 3
    
    return 0


def process_paragraph(paragraph, text, base_font_size, paragraph_counter, h1_used):
    """
    Обрабатывает параграф и возвращает соответствующий HTML элемент
    
    Args:
        paragraph: Объект параграфа python-docx
        text (str): Очищенный текст
        base_font_size (float): Базовый размер шрифта
        paragraph_counter (int): Номер абзаца
        h1_used (bool): Использован ли уже заголовок h1
        
    Returns:
        str: HTML элемент
    """
    heading_level = is_heading(paragraph, text, base_font_size)
    
    if heading_level > 0:
        # Если это первый заголовок и h1 еще не использован
        if heading_level == 1 and not h1_used:
            return f'<h1>{text}</h1>'
        elif heading_level == 1:
            # Если h1 уже использован, делаем h2
            return f'<h2>{text}</h2>'
        elif heading_level == 2:
            return f'<h2>{text}</h2>'
        else:
            return f'<h3>{text}</h3>'
    else:
        # Обычный абзац с номером
        return f'<p{paragraph_counter:02d}>{text}</p{paragraph_counter:02d}>'


def extract_tables_as_html(document):
    """
    Извлекает таблицы из документа и конвертирует в HTML
    
    Args:
        document: Объект документа python-docx
        
    Returns:
        str: HTML код таблиц
    """
    if not document.tables:
        return ""
    
    html_tables = []
    
    for table_idx, table in enumerate(document.tables):
        rows_html = []
        
        for row_idx, row in enumerate(table.rows):
            cells_html = []
            
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cleaned_cell = clean_text(cell_text)
                    # Для первой строки используем th (заголовки), для остальных td
                    tag = 'th' if row_idx == 0 else 'td'
                    cells_html.append(f'<{tag}>{cleaned_cell}</{tag}>')
                else:
                    tag = 'th' if row_idx == 0 else 'td'
                    cells_html.append(f'<{tag}></{tag}>')
            
            if cells_html:
                rows_html.append(f'<tr>{"".join(cells_html)}</tr>')
        
        if rows_html:
            table_html = f'<table>{"".join(rows_html)}</table>'
            html_tables.append(table_html)
    
    return ''.join(html_tables)


def clean_text(text):
    """
    Очищает текст от спецсимволов, оставляя только читаемые символы
    
    Args:
        text (str): Исходный текст
        
    Returns:
        str: Очищенный текст
    """
    # Убираем управляющие символы, кроме переносов строк и табуляций
    cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    # Нормализуем пробелы - заменяем множественные пробелы на одинарные
    cleaned = re.sub(r'\s+', ' ', cleaned)
    
    # Убираем пробелы в начале и конце
    cleaned = cleaned.strip()
    
    # Экранируем HTML символы
    cleaned = cleaned.replace('&', '&amp;')
    cleaned = cleaned.replace('<', '&lt;')
    cleaned = cleaned.replace('>', '&gt;')
    cleaned = cleaned.replace('"', '&quot;')
    cleaned = cleaned.replace("'", '&#39;')
    
    return cleaned


def get_document_structure(file_content):
    """
    Анализирует структуру документа и возвращает статистику
    
    Args:
        file_content (bytes): Содержимое DOCX файла в байтах
        
    Returns:
        dict: Статистика структуры документа
    """
    try:
        file_stream = BytesIO(file_content)
        document = Document(file_stream)
        
        font_sizes = analyze_font_sizes(document)
        
        paragraphs_count = 0
        headings_count = 0
        tables_count = len(document.tables)
        
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                if is_heading(paragraph, text, font_sizes['base']):
                    headings_count += 1
                else:
                    paragraphs_count += 1
        
        return {
            'paragraphs': paragraphs_count,
            'headings': headings_count,
            'tables': tables_count,
            'font_sizes': font_sizes
        }
        
    except Exception as e:
        return {'error': str(e)}